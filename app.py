import streamlit as st
import google.generativeai as genai
from docxtpl import DocxTemplate
from docx import Document
import io
import json
import pandas as pd
from datetime import datetime
import os
from PIL import Image
import re
import traceback
import time
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception, before_sleep

# ==================================================
# 🔧 RUTAS (ARCHIVOS EN EL MISMO DIRECTORIO - Streamlit Cloud)
# ==================================================
BASE_DIR = "."  # directorio actual donde está la app
EXCEL_PATH = "Listas Maestras de Documentos.xlsx"
LOGO_PATH = "logo_daka.png"
PLANTILLA_PATH = "Formato de Procedimiento.docx"

# ==================================================
# 🤖 CONFIGURACIÓN DE GEMINI
# ==================================================
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    st.error("❌ No se encontró la API key. Configúrala en los 'Secrets' de Streamlit Cloud con el nombre 'GEMINI_API_KEY'.")
    st.stop()
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

# ==================================================
# 📥 FUNCIONES DE CARGA DE DATOS
# ==================================================
@st.cache_data
def cargar_niveles():
    try:
        df = pd.read_excel(EXCEL_PATH, sheet_name="Niveles Operativos")
        nombres = df['Empresas'].astype(str).str.strip()
        abreviaturas = df['Abreviatura'].astype(str).str.strip()
        # Filtrar aquellos donde el nombre no sea 'nan' ni vacío
        niveles = {}
        for i in range(len(df)):
            nombre = nombres.iloc[i]
            abrev = abreviaturas.iloc[i]
            if nombre and nombre != 'nan' and abrev and abrev != 'nan':
                niveles[nombre] = abrev
        return niveles
    except Exception as e:
        st.error(f"Error cargando Niveles: {e}")
        # Fallback manual (sin 'nan')
        return {"Corporativo":"CO","Tienda":"TD","Centro de Distribución":"CD","Logi Express":"LE","Smartech":"ST"}

@st.cache_data
def cargar_tipos():
    try:
        df = pd.read_excel(EXCEL_PATH, sheet_name="Tipos de Documento")
        nombres = df['Tipo de Documento'].astype(str).str.strip()
        abreviaturas = df['Abreviatura'].astype(str).str.strip()
        tipos = {nombres[i]: abreviaturas[i] for i in range(len(df))}
        return tipos
    except Exception as e:
        st.error(f"Error cargando Tipos: {e}")
        return {"Procedimiento":"PR","Instrucción de Trabajo":"IT","Política":"PO"}

@st.cache_data
def cargar_procesos():
    try:
        df = pd.read_excel(EXCEL_PATH, sheet_name="Proceso")
        nivel_col = 'Nivel Operativo'
        abrev_col = 'Abreviatura'
        nombre_col = 'Nombre del Proceso y Sub Proceso'
        procesos = {}
        for _, row in df.iterrows():
            nivel = str(row[nivel_col]).strip()
            nombre = str(row[nombre_col]).strip()
            abrev = str(row[abrev_col]).strip()
            if nivel not in procesos:
                procesos[nivel] = []
            procesos[nivel].append((nombre, abrev))
        return procesos
    except Exception as e:
        st.error(f"Error cargando Procesos: {e}")
        return {}

@st.cache_data
def cargar_codigos_existentes():
    try:
        df = pd.read_excel(EXCEL_PATH, sheet_name="Lista Maestra Interna")
        if 'Código' in df.columns:
            codigos = df['Código'].dropna().astype(str).tolist()
        else:
            codigos = df.iloc[:, 3].dropna().astype(str).tolist()
        return codigos
    except Exception:
        return []

def calcular_correlativo(nivel_abrev, tipo_abrev, proceso_abrev, codigos):
    patron = f"{nivel_abrev}-{tipo_abrev}-{proceso_abrev}-"
    numeros = []
    for c in codigos:
        if c.startswith(patron):
            parte = c.split("-")[-1]
            if parte.isdigit():
                numeros.append(int(parte))
    if numeros:
        return f"{max(numeros)+1:03d}"
    return "001"

# ==================================================
# 📝 PROMPT DEFINITIVO (con formato de numeración exacto)
# ==================================================
PROMPT_GEMINI = """
Eres un analista de calidad con 10 años de experiencia redactando procedimientos bajo ISO 9001:2015. Tu tarea es analizar la transcripción de una reunión de trabajo y generar un JSON con el contenido completo de un procedimiento.

Debes seguir ESTRICTAMENTE las siguientes reglas de formato:

## 1. Título del documento
- Genera un título profesional en mayúsculas sostenidas que refleje el tema central del procedimiento.
- Comienza con "PROCEDIMIENTO DE ..." (ej: "PROCEDIMIENTO DE GESTIÓN DE PRODUCTOS NO CONFORMES").

## 2. Objetivo
- Redacta una oración que comience con "Establecer..." o "Definir...".

## 3. Alcance
- Indica los límites del procedimiento (áreas, procesos o roles incluidos/excluidos).

## 4. Definiciones
- Lista en viñetas (con guiones) cada término y su definición.

## 5. Responsabilidades
- Lista en viñetas (con guiones) cada cargo o área y su responsabilidad.

## 6. Pasos (tabla PASO A PASO)
Genera un array llamado "pasos". Cada elemento es un objeto con los siguientes campos:
- "numero_paso": número entero (1, 2, 3, ...). Corresponde al número principal de la etapa.
- "subactividad": string con formato "X.Y" (ej: "1.1", "1.2", "2.1").
- "etapa": nombre corto de la etapa (ej: "Recepción", "Inspección").
- "actividad": descripción detallada de la acción, SIN incluir el número de subactividad.
- "responsable": cargo o área responsable.

**CRÍTICO**: 
- El campo "numero_paso" debe ser un número entero (1, 2, 3...), NO un decimal.
- El campo "subactividad" debe ser string con formato "X.Y".
- Puedes tener múltiples subactividades para el mismo paso principal (ej: 1.1, 1.2, 1.3...). Para ello, crea un objeto por cada subactividad, con el mismo "numero_paso" y diferente "subactividad".

Ejemplo de un paso principal con dos subactividades:
[
  {"numero_paso": 1, "subactividad": "1.1", "etapa": "Recepción", "actividad": "Recibir la guía de despacho del transportista.", "responsable": "Auxiliar de Almacén"},
  {"numero_paso": 1, "subactividad": "1.2", "etapa": "Recepción", "actividad": "Verificar la cantidad de bultos o unidades declaradas en la guía.", "responsable": "Auxiliar de Almacén"}
]

## 7. Documentos asociados
- Array de objetos con "nombre" y "codigo". Si no se menciona ninguno, incluye al menos un documento por defecto.

## 8. Formato JSON de salida
Devuelve ÚNICAMENTE el JSON, sin texto adicional. La estructura exacta es:

{
  "titulo": "PROCEDIMIENTO DE ...",
  "objetivo": "texto",
  "alcance": "texto",
  "definiciones": "- término: definición\\n- otro: definición",
  "responsabilidades": "- cargo: actividad\\n- otro: actividad",
  "pasos": [...],
  "documentos": [{"nombre": "...", "codigo": "..."}]
}

Transcripción de la reunión:
"""

# ==================================================
# 🔁 LLAMADA A GEMINI CON REINTENTOS
# ==================================================
def is_quota_error(exception):
    return isinstance(exception, genai.types.generation_types.StopCandidateException) and "429" in str(exception)

def log_retry(retry_state):
    msg = f"⚠️ Cuota excedida (429). Reintentando en {retry_state.next_action.sleep:.0f}s... (intento {retry_state.attempt_number})"
    st.warning(msg)
    print(msg)

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=2, min=4, max=30),
    retry=retry_if_exception(is_quota_error),
    before_sleep=log_retry,
    reraise=True
)
def _call_gemini(texto):
    respuesta = model.generate_content(PROMPT_GEMINI + texto)
    raw = respuesta.text
    raw = re.sub(r"```json\s*|\s*```", "", raw.strip())
    return json.loads(raw)

def extraer_json(texto):
    try:
        return _call_gemini(texto)
    except Exception as e:
        if "429" in str(e):
            st.error("❌ Límite de cuota excedido incluso después de reintentos. Espera unos minutos o cambia a plan de pago.")
        else:
            st.error(f"Error en Gemini: {e}")
        return {}

# ==================================================
# 🎨 INTERFAZ DE USUARIO
# ==================================================
st.set_page_config(page_title="SGC Daka - Generador de Procedimientos", layout="wide")

if os.path.exists(LOGO_PATH):
    col1, col2 = st.columns([1,5])
    with col1:
        st.image(Image.open(LOGO_PATH), width=80)
    with col2:
        st.title("🏛️ Generador Automático de Procedimientos SGC Daka")
else:
    st.title("🏛️ Generador Automático de Procedimientos SGC Daka")
    st.caption("(Logo no encontrado)")

with st.sidebar:
    if os.path.exists(LOGO_PATH):
        st.image(Image.open(LOGO_PATH), use_container_width=True)
    st.header("⚙️ Control Documental")
    
    niveles_dict = cargar_niveles()
    tipos_dict = cargar_tipos()
    procesos_por_nivel = cargar_procesos()
    codigos = cargar_codigos_existentes()
    
    nivel_nombre = st.selectbox("Nivel Operativo", list(niveles_dict.keys()))
    nivel_abrev = niveles_dict[nivel_nombre]
    
    tipo_nombre = st.selectbox("Tipo de Documento", list(tipos_dict.keys()))
    tipo_abrev = tipos_dict[tipo_nombre]
    
    procesos_disponibles = procesos_por_nivel.get(nivel_abrev, [])
    if not procesos_disponibles:
        st.warning(f"No hay procesos definidos para el nivel {nivel_nombre} (abreviatura: {nivel_abrev})")
        proceso_abrev = "XX"
        proceso_nombre = "Sin procesos"
    else:
        opciones = [f"{nombre} ({abrev})" for nombre, abrev in procesos_disponibles]
        proceso_seleccionado = st.selectbox("Proceso", opciones)
        proceso_nombre = proceso_seleccionado.split(" (")[0]
        proceso_abrev = proceso_seleccionado.split("(")[-1].strip(")")
    
    correlativo_sugerido = calcular_correlativo(nivel_abrev, tipo_abrev, proceso_abrev, codigos)
    correlativo = st.text_input("Correlativo", value=correlativo_sugerido)
    if not correlativo.isdigit():
        st.error("El correlativo debe ser numérico")
        correlativo = "001"
    
    codigo_final = f"{nivel_abrev}-{tipo_abrev}-{proceso_abrev}-{correlativo}"
    st.success(f"📌 Código: `{codigo_final}`")

st.markdown("---")
st.subheader("📄 Transcripción de la reunión (TurboScribe)")
entrevista = st.text_area("Pega aquí el texto:", height=300)

# ==================================================
# 📊 FUNCIONES PARA RELLENAR TABLAS (CORREGIDAS)
# ==================================================
def llenar_tabla_pasos(doc, pasos):
    """
    Normaliza y llena la tabla de PASO A PASO con el formato deseado:
    - Columna N°: número entero (1, 2, 3...)
    - Columna ACTIVIDAD: subactividad + " " + actividad (ej: "1.1 Recibir la guía...")
    """
    # Normalizar pasos (asegurar que tengan numero_paso (int) y subactividad (str))
    pasos_norm = []
    for idx, p in enumerate(pasos):
        # Determinar numero_paso (entero)
        if 'numero_paso' in p:
            np = p['numero_paso']
            if isinstance(np, str) and '.' in np:
                np = int(np.split('.')[0])
            elif isinstance(np, float):
                np = int(np)
            elif isinstance(np, str) and np.isdigit():
                np = int(np)
            else:
                np = idx+1
        elif 'numero' in p:
            raw = str(p['numero'])
            if '.' in raw:
                np = int(raw.split('.')[0])
            else:
                np = int(raw) if raw.isdigit() else idx+1
        else:
            np = idx+1

        # Determinar subactividad
        if 'subactividad' in p and p['subactividad']:
            sub = str(p['subactividad'])
        elif 'numero' in p:
            sub = str(p['numero'])
        else:
            # Generar automática
            count = sum(1 for x in pasos_norm if x['numero_paso'] == np) + 1
            sub = f"{np}.{count}"

        # Asegurar formato correcto: si sub no tiene punto o no empieza con np, reajustar
        if '.' not in sub:
            count = sum(1 for x in pasos_norm if x['numero_paso'] == np) + 1
            sub = f"{np}.{count}"
        else:
            parte_principal = int(sub.split('.')[0])
            if parte_principal != np:
                count = sum(1 for x in pasos_norm if x['numero_paso'] == np) + 1
                sub = f"{np}.{count}"

        paso_ok = {
            'numero_paso': np,
            'subactividad': sub,
            'etapa': p.get('etapa', ''),
            'actividad': p.get('actividad', ''),
            'responsable': p.get('responsable', '')
        }
        pasos_norm.append(paso_ok)

    # Llenar la tabla en el documento
    for tabla in doc.tables:
        if len(tabla.rows) and len(tabla.columns) >= 4 and "N°" in tabla.rows[0].cells[0].text:
            for i, p in enumerate(pasos_norm):
                if i + 1 >= len(tabla.rows):
                    tabla.add_row()
                row = tabla.rows[i + 1]
                row.cells[0].text = str(p['numero_paso'])                    # N° entero
                row.cells[1].text = p['etapa']
                row.cells[2].text = f"{p['subactividad']} {p['actividad']}"  # ACTIVIDAD con subnumeración
                row.cells[3].text = p['responsable']
            break

def llenar_tabla_documentos(doc, docs):
    for tabla in doc.tables:
        if len(tabla.rows) and "DOCUMENTOS ASOCIADOS" in tabla.rows[0].cells[0].text:
            if len(tabla.rows) >= 2:
                celda_nom = tabla.rows[1].cells[0]
                celda_cod = tabla.rows[1].cells[1]
                if docs:
                    texto_nom = "\n".join(f"- {d['nombre']}" for d in docs)
                    texto_cod = "\n".join(d.get("codigo", "") for d in docs)
                else:
                    texto_nom = "Ninguno"
                    texto_cod = ""
                celda_nom.text = texto_nom
                celda_cod.text = texto_cod
            break

def llenar_resumen_cambios(doc, version, resumen, fecha):
    for tabla in doc.tables:
        if len(tabla.rows) and "N° VERSIÓN" in tabla.rows[0].cells[0].text:
            if len(tabla.rows) >= 2:
                fila = tabla.rows[1]
                fila.cells[0].text = version
                fila.cells[1].text = resumen
                fila.cells[2].text = fecha
            break

# ==================================================
# 🚀 BOTÓN GENERAR PROCEDIMIENTO
# ==================================================
if st.button("🚀 Generar Procedimiento", type="primary"):
    if not entrevista.strip():
        st.error("❌ Transcripción vacía")
    else:
        with st.spinner("Analizando con Gemini (puede tardar 10-20 segundos)..."):
            datos = extraer_json(entrevista)
            if not datos:
                st.stop()
        with st.expander("📋 Datos extraídos por Gemini (revisa antes de generar)"):
            st.json(datos)
        
        # Asegurar que haya al menos un documento asociado
        documentos = datos.get("documentos", [])
        if not documentos:
            documentos = [{"nombre": "Formato de Registro de Gestión", "codigo": "F-SGC-001"}]
        
        # Título: usar el generado por Gemini o fallback
        titulo_final = datos.get("titulo", "")
        if not titulo_final:
            titulo_final = f"{tipo_nombre.upper()} DE {proceso_nombre.upper()}"
        
        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        version = "00"
        resumen_cambio = "Nuevo Documento"
        
        contexto = {
            "titulo": titulo_final,
            "proceso": proceso_nombre,
            "codigo": codigo_final,
            "version": version,
            "fecha_emision": fecha_hoy,
            "resumen_cambio": resumen_cambio,
            "fecha_aprobacion": fecha_hoy,
            "objetivo": datos.get("objetivo", ""),
            "alcance": datos.get("alcance", ""),
            "definiciones": datos.get("definiciones", ""),
            "responsabilidades": datos.get("responsabilidades", "")
        }
        
        if not os.path.exists(PLANTILLA_PATH):
            st.error(f"❌ Plantilla no encontrada: {PLANTILLA_PATH}")
            st.stop()
        
        try:
            doc_tpl = DocxTemplate(PLANTILLA_PATH)
            doc_tpl.render(contexto)
            buf = io.BytesIO()
            doc_tpl.save(buf)
            buf.seek(0)
            doc_final = Document(buf)
            llenar_tabla_pasos(doc_final, datos.get("pasos", []))
            llenar_tabla_documentos(doc_final, documentos)
            llenar_resumen_cambios(doc_final, version, resumen_cambio, fecha_hoy)
            out = io.BytesIO()
            doc_final.save(out)
            out.seek(0)
            st.success("✅ Procedimiento generado correctamente")
            st.download_button(
                label="📥 Descargar DOCX",
                data=out,
                file_name=f"{codigo_final}_Procedimiento_v{version}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Error al generar: {e}")
            st.code(traceback.format_exc())

print("✅ app.py cargado correctamente con todas las correcciones (numeración entera + subactividades)")
